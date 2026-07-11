"""Turn an input (uploaded file or pasted text) into a list of chapters.

Each chapter is a dict: {"title": str, "text": str}.
Supported: .txt, .md, .epub, .pdf, .docx, and raw text from the editor.
"""
import os
import re


def _clean(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # collapse runs of blank lines, strip trailing spaces
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _split_markdown(text: str, default_title: str):
    """Split on markdown H1/H2 headings into chapters; otherwise one chapter."""
    lines = text.split("\n")
    chapters = []
    cur_title, cur_lines = None, []

    def flush():
        body = _clean("\n".join(cur_lines))
        if body:
            chapters.append({"title": cur_title or default_title, "text": body,
                             "heading": cur_title is not None})

    for ln in lines:
        m = re.match(r"^#{1,6}\s+(.*)$", ln.strip())
        if m:
            flush()
            cur_title, cur_lines = m.group(1).strip(), []
        else:
            cur_lines.append(ln)
    flush()
    if not chapters:
        chapters = [{"title": default_title, "text": _clean(text), "heading": False}]
    return chapters


def from_text(text: str, title: str = "Story"):
    return _split_markdown(text, title)


def from_txt(path: str):
    title = os.path.splitext(os.path.basename(path))[0]
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return _split_markdown(f.read(), title)


def from_pdf(path: str):
    from pypdf import PdfReader
    reader = PdfReader(path)
    title = (reader.metadata.title if reader.metadata else None) or \
        os.path.splitext(os.path.basename(path))[0]
    pages = [(p.extract_text() or "") for p in reader.pages]

    # Split on the PDF outline (bookmarks) when present, so a book becomes real
    # chapters instead of one giant blob.
    marks = []
    try:
        def walk(outline):
            for it in outline:
                if isinstance(it, list):
                    walk(it)
                    continue
                try:
                    pnum = reader.get_destination_page_number(it)
                except Exception:  # noqa: BLE001
                    pnum = None
                ttl = (getattr(it, "title", None) or "").strip()
                if pnum is not None and ttl:
                    marks.append((pnum, ttl))
        walk(reader.outline)
    except Exception:  # noqa: BLE001
        marks = []

    chapters = []
    bounds = []
    for pnum, ttl in sorted(marks):
        if bounds and bounds[-1][0] == pnum:   # one title per page boundary
            continue
        bounds.append((pnum, ttl))
    if len(bounds) >= 2:
        for i, (pstart, ttl) in enumerate(bounds):
            pend = bounds[i + 1][0] if i + 1 < len(bounds) else len(pages)
            body = _clean("\n\n".join(pages[pstart:pend]))
            if body:
                chapters.append({"title": ttl, "text": body, "heading": True})
        if bounds[0][0] > 0:                   # front matter before the first mark
            front = _clean("\n\n".join(pages[:bounds[0][0]]))
            if front:
                chapters.insert(0, {"title": title, "text": front, "heading": False})

    if not chapters:                           # no usable outline → single chapter
        text = _clean("\n\n".join(pages))
        chapters = [{"title": title, "text": text, "heading": False}] if text else []
    return chapters


def from_docx(path: str):
    import docx
    doc = docx.Document(path)
    title = os.path.splitext(os.path.basename(path))[0]
    chapters = []
    cur_title, cur_lines = None, []

    def flush():
        body = _clean("\n".join(cur_lines))
        if body:
            chapters.append({"title": cur_title or title, "text": body,
                             "heading": cur_title is not None})

    for p in doc.paragraphs:
        style = (p.style.name or "").lower() if p.style else ""
        if style.startswith("heading 1") or style.startswith("title"):
            flush()
            cur_title, cur_lines = p.text.strip(), []
        else:
            cur_lines.append(p.text)
    flush()
    return chapters or [{"title": title, "heading": False,
                         "text": _clean("\n".join(p.text for p in doc.paragraphs))}]


def _epub_toc_map(book):
    """Map each content file (by basename) to its real chapter title from the
    book's TOC/nav, so we name chapters the way the book intends instead of
    guessing from the first heading. Outermost entry wins for a given file."""
    from ebooklib import epub
    mapping = {}

    def add(href, title):
        base = os.path.basename((href or "").split("#")[0])
        title = (title or "").strip()
        if base and title and base not in mapping:
            mapping[base] = title

    def walk(items):
        for it in items or []:
            if isinstance(it, (tuple, list)):
                node = it[0] if it else None
                children = it[1] if len(it) > 1 else []
                if node is not None:
                    walk([node])
                walk(children)
            elif isinstance(it, epub.Link):
                add(it.href, it.title)
            elif isinstance(it, epub.Section):
                add(getattr(it, "href", ""), getattr(it, "title", ""))

    walk(getattr(book, "toc", []))
    return mapping


def from_epub(path: str):
    from ebooklib import epub, ITEM_DOCUMENT
    from bs4 import BeautifulSoup
    book = epub.read_epub(path)
    title = (book.get_metadata("DC", "title") or [["Audiobook"]])[0][0]
    toc_map = _epub_toc_map(book)

    # Walk the spine (the publisher's reading order) rather than manifest order,
    # so chapters come out ordered and complete.
    items = []
    for entry in (book.spine or []):
        idref = entry[0] if isinstance(entry, (tuple, list)) else entry
        it = book.get_item_with_id(idref)
        if it is not None:
            items.append(it)
    if not items:                        # fall back to manifest order
        items = list(book.get_items_of_type(ITEM_DOCUMENT))

    chapters = []
    for item in items:
        if item.get_type() != ITEM_DOCUMENT:
            continue
        soup = BeautifulSoup(item.get_content(), "html.parser")
        for tag in soup(["script", "style", "nav"]):
            tag.decompose()
        text = _clean(soup.get_text("\n"))
        if len(text) < 20:               # skip covers, nav, empty pages
            continue
        # prefer the real TOC title, then a heading, then a running number
        ch_title = toc_map.get(os.path.basename((item.get_name() or "").split("#")[0]))
        from_heading = ch_title is not None
        if not ch_title:
            heading = soup.find(["h1", "h2", "h3"])
            ch_title = heading.get_text(strip=True) if heading else None
            from_heading = ch_title is not None
        if not ch_title:
            ch_title = f"Chapter {len(chapters) + 1}"
        chapters.append({"title": ch_title, "text": text, "heading": from_heading})
    return chapters, title


def from_html(path: str):
    from bs4 import BeautifulSoup
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        soup = BeautifulSoup(f.read(), "html.parser")
    for tag in soup(["script", "style", "nav"]):
        tag.decompose()
    title = (soup.title.get_text(strip=True) if soup.title else "") \
        or os.path.splitext(os.path.basename(path))[0]
    text = _clean(soup.get_text("\n"))
    return [{"title": title, "text": text, "heading": False}] if text else []


def from_rtf(path: str):
    from striprtf.striprtf import rtf_to_text
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        text = rtf_to_text(f.read())
    return _split_markdown(text, os.path.splitext(os.path.basename(path))[0])


def from_fb2(path: str):
    """FictionBook2 (XML). One chapter per leaf <section>, titled from its
    <title>. Parsed as real XML (FB2's <title> holds <p> children, which an
    HTML parser would mangle) and namespace-agnostic. Notes bodies are skipped."""
    from lxml import etree
    base = os.path.splitext(os.path.basename(path))[0]
    try:
        root = etree.parse(path).getroot()
    except Exception:  # noqa: BLE001
        return []

    def lname(el):
        try:
            return etree.QName(el).localname
        except Exception:  # noqa: BLE001 — comments/PIs
            return ""

    chapters = []

    def emit(sec):
        subs = [c for c in sec if lname(c) == "section"]
        if subs:                         # recurse into nested sections
            for s in subs:
                emit(s)
            return
        ttl = None
        for c in sec:
            if lname(c) == "title":
                ttl = re.sub(r"\s+", " ", " ".join(c.itertext())).strip() or None
                break
        text = _clean("\n".join(" ".join(p.itertext())
                                for p in sec if lname(p) == "p"))
        if len(text) >= 20:
            chapters.append({"title": ttl or f"Chapter {len(chapters) + 1}",
                            "text": text, "heading": ttl is not None})

    bodies = [e for e in root if lname(e) == "body"]
    for body in bodies:
        if body.get("name"):             # skip <body name="notes"> footnote bodies
            continue
        for sec in body:
            if lname(sec) == "section":
                emit(sec)
    if not chapters and bodies:
        text = _clean("\n".join(bodies[0].itertext()))
        if text:
            chapters = [{"title": base, "text": text, "heading": False}]
    return chapters


EXT_HANDLERS = {
    ".txt": from_txt, ".md": from_txt, ".markdown": from_txt,
    ".pdf": from_pdf, ".docx": from_docx,
    ".html": from_html, ".htm": from_html,
    ".rtf": from_rtf, ".fb2": from_fb2,
}
# .epub and .kepub (Kobo) are handled directly (they return their own title).
SUPPORTED_EXTS = set(EXT_HANDLERS) | {".epub", ".kepub"}


def extract(path: str):
    """Returns (chapters, book_title)."""
    ext = os.path.splitext(path)[1].lower()
    base = os.path.splitext(os.path.basename(path))[0]
    if ext in (".epub", ".kepub"):
        return from_epub(path)
    handler = EXT_HANDLERS.get(ext)
    if not handler:
        raise ValueError(f"Unsupported file type: {ext}")
    return handler(path), base
